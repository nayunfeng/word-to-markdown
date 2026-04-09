from __future__ import annotations

import base64
import subprocess
import tempfile
import unittest
from unittest.mock import patch
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from word_to_markdown.conversion import ConversionError, prepare_input
from word_to_markdown.parser import parse_docx
from word_to_markdown.processor import process_batch

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9n0n4AAAAASUVORK5CYII="
)


class EndToEndTests(unittest.TestCase):
    def test_process_docx_outputs_markdown_sections_and_media(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "sample.docx"
            image = root / "tiny.png"
            output = root / "output"
            image.write_bytes(PNG_BYTES)

            document = Document()
            document.add_heading("项目背景", level=1)
            document.add_paragraph("这是第一段。")
            document.add_paragraph("- 列表项目")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "列1"
            table.cell(0, 1).text = "列2"
            table.cell(1, 0).text = "值A"
            table.cell(1, 1).text = "值B"
            document.add_picture(str(image))
            document.save(source)

            batch, warnings = process_batch(
                input_path=source,
                output_root=output,
                recursive=False,
                split=True,
                split_level=1,
                extract_images=True,
                doc_converter="auto",
            )

            self.assertEqual([], warnings)
            self.assertEqual(1, batch.total)
            self.assertEqual(1, batch.success_count)

            result = batch.results[0]
            section_md = (result.output_dir / "项目背景" / "项目背景.md").read_text(encoding="utf-8")

            self.assertFalse((result.output_dir / "full.md").exists())
            self.assertFalse((result.output_dir / "sections").exists())
            self.assertIn("# 项目背景", section_md)
            self.assertIn("- 列表项目", section_md)
            self.assertIn("| 列1 | 列2 |", section_md)
            self.assertIn("![image_001.png](../media/image_001.png)", section_md)

    def test_batch_partial_result_when_split_heading_missing(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docs_dir = root / "docs"
            docs_dir.mkdir()
            output = root / "output"
            source = docs_dir / "plain.docx"

            document = Document()
            document.add_paragraph("没有标题，只有正文。")
            document.save(source)

            batch, warnings = process_batch(
                input_path=docs_dir,
                output_root=output,
                recursive=False,
                split=True,
                split_level=1,
                extract_images=False,
                doc_converter="auto",
            )

            self.assertEqual([], warnings)
            self.assertEqual(1, batch.total)
            self.assertEqual(1, batch.success_count)
            self.assertTrue((batch.results[0].output_dir / "文档概览.md").exists())

    def test_doc_input_raises_clear_error_without_converter(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "legacy.doc"
            source.write_bytes(b"fake-doc-content")

            with patch("word_to_markdown.conversion.select_converter", return_value=None):
                with self.assertRaises(ConversionError) as context:
                    prepare_input(source, converter="auto")

            self.assertIn("No available converter found", str(context.exception))

    def test_word_converter_timeout_raises_clear_error(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "legacy.doc"
            source.write_bytes(b"fake-doc-content")

            with patch("word_to_markdown.conversion.shutil.which", return_value="powershell"):
                with patch(
                    "word_to_markdown.conversion.subprocess.run",
                    side_effect=subprocess.TimeoutExpired(cmd="powershell", timeout=60),
                ):
                    with self.assertRaises(ConversionError) as context:
                        prepare_input(source, converter="word")

            self.assertIn("Microsoft Word automation timed out", str(context.exception))

    def test_parser_skips_toc_and_detects_custom_heading_style(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "toc.docx"

            document = Document()
            styles = document.styles
            toc_style = styles.add_style("toc 1", WD_STYLE_TYPE.PARAGRAPH)
            custom_heading = styles.add_style("样式 标题 1H1 heading 1", WD_STYLE_TYPE.PARAGRAPH)

            document.add_paragraph("目 录")
            paragraph = document.add_paragraph("第一章 报文结构\t1")
            paragraph.style = toc_style
            paragraph = document.add_paragraph("报文结构")
            paragraph.style = custom_heading
            document.add_paragraph("正文内容")
            document.save(source)

            parsed = parse_docx(source=source, media_dir=root / "media", extract_images=False)
            kinds = [block.kind for block in parsed.blocks]
            texts = [getattr(block, "text", "") for block in parsed.blocks]

            self.assertEqual(["heading", "paragraph"], kinds)
            self.assertEqual("报文结构", texts[0])
            self.assertEqual("正文内容", texts[1])

    def test_xml_paragraphs_render_as_code_block(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "xml.docx"
            output = root / "output"

            document = Document()
            document.add_heading("接口示例", level=1)
            document.add_paragraph("请求报文如下：")
            document.add_paragraph('<?xml version="1.0" encoding="GBK"?>')
            document.add_paragraph("<stream>")
            document.add_paragraph("<action>DLMACSUB</action>")
            document.add_paragraph("</stream>")
            document.add_paragraph("以上为示例。")
            document.save(source)

            batch, warnings = process_batch(
                input_path=source,
                output_root=output,
                recursive=False,
                split=True,
                split_level=1,
                extract_images=False,
                doc_converter="auto",
            )

            self.assertEqual([], warnings)
            result = batch.results[0]
            section_md = (result.output_dir / "接口示例" / "接口示例.md").read_text(encoding="utf-8")

            self.assertFalse((result.output_dir / "文档概览.md").exists())
            self.assertIn("```", section_md)
            self.assertIn("```xml", section_md)
            self.assertIn("<action>DLMACSUB</action>", section_md)
            self.assertIn("以上为示例。", section_md)

    def test_split_exports_hierarchical_tree(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "tree.docx"
            output = root / "output"

            document = Document()
            document.add_heading("第一章 总览", level=1)
            document.add_heading("1.1 背景", level=2)
            document.add_paragraph("背景内容")
            document.add_heading("第二章 接口", level=1)
            document.add_heading("2.1 薪资代发", level=2)
            document.add_paragraph("代发分组说明")
            document.add_heading("2.1.1 代发提交", level=3)
            document.add_paragraph("提交内容")
            document.add_heading("2.1.1.1 请求报文", level=4)
            document.add_paragraph("请求细节")
            document.add_heading("2.1.2 结果查询", level=3)
            document.add_paragraph("查询内容")
            document.save(source)

            batch, warnings = process_batch(
                input_path=source,
                output_root=output,
                recursive=False,
                split=True,
                split_level=1,
                extract_images=False,
                doc_converter="auto",
            )

            self.assertEqual([], warnings)
            result = batch.results[0]
            chapter_one = result.output_dir / "第一章 总览" / "第一章 总览.md"
            group_readme = result.output_dir / "第二章 接口" / "2.1 薪资代发" / "2.1 薪资代发.md"
            leaf_file = result.output_dir / "第二章 接口" / "2.1 薪资代发" / "2.1.1 代发提交.md"

            self.assertTrue(chapter_one.exists())
            self.assertTrue(group_readme.exists())
            self.assertTrue(leaf_file.exists())
            self.assertIn("## 1.1 背景", chapter_one.read_text(encoding="utf-8"))
            self.assertIn("## 2.1 薪资代发", group_readme.read_text(encoding="utf-8"))
            leaf_text = leaf_file.read_text(encoding="utf-8")
            self.assertIn("### 2.1.1 代发提交", leaf_text)
            self.assertIn("#### 2.1.1.1 请求报文", leaf_text)

    def test_split_many_second_level_children_into_multiple_files(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "appendix.docx"
            output = root / "output"

            document = Document()
            document.add_heading("附录", level=1)
            document.add_heading("4.1 证件类型", level=2)
            document.add_paragraph("证件类型内容")
            document.add_heading("4.2 制单状态", level=2)
            document.add_paragraph("制单状态内容")
            document.add_heading("4.3 交易类型", level=2)
            document.add_paragraph("交易类型内容")
            document.save(source)

            batch, warnings = process_batch(
                input_path=source,
                output_root=output,
                recursive=False,
                split=True,
                split_level=1,
                extract_images=False,
                doc_converter="auto",
            )

            self.assertEqual([], warnings)
            result = batch.results[0]
            appendix_dir = result.output_dir / "附录"
            self.assertTrue((appendix_dir / "4.1 证件类型.md").exists())
            self.assertTrue((appendix_dir / "4.2 制单状态.md").exists())
            self.assertTrue((appendix_dir / "4.3 交易类型.md").exists())
            self.assertFalse((appendix_dir / "附录.md").exists())


if __name__ == "__main__":
    unittest.main()
