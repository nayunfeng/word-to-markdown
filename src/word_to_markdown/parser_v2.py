from docx import Document


def parse_docx(path):
    doc = Document(path)
    elements = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = para.style.name.lower()

        if 'heading' in style:
            try:
                level = int(style.replace('heading', '').strip() or 1)
            except:
                level = 1
            elements.append({'type': 'heading', 'level': level, 'text': text})
            continue

        if 'list' in style or text.startswith(('•', '-', '*')):
            elements.append({'type': 'list', 'text': text})
            continue

        elements.append({'type': 'paragraph', 'text': text})

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)

        elements.append({'type': 'table', 'data': table_data})

    return elements
